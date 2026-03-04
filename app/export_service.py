"""
Export Service for Business Strategy Reports
Generates DOCX and PDF files from strategy report data.
"""

import io
import json
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT


class StrategyReportExporter:
    """Export strategy reports to DOCX and PDF formats."""
    
    def __init__(self, report_data: Dict[str, Any]):
        """
        Initialize exporter with report data.
        
        Args:
            report_data: Dictionary containing strategy report fields
        """
        self.data = report_data
        self.company_name = report_data.get('company_name', '')
        self.location = report_data.get('location', '')
        self.usp = report_data.get('usp', '')
        self.business_type = report_data.get('business_type', 'Business')
        self.target_audience = report_data.get('target_audience', 'General')
        self.budget = report_data.get('budget', 'Not specified')
        self.created_at = report_data.get('created_at', datetime.now().isoformat())
        
        # Parse full_report if it's a JSON string
        if isinstance(report_data.get('full_report'), str):
            try:
                self.full_report = json.loads(report_data['full_report'])
            except (json.JSONDecodeError, TypeError):
                self.full_report = {}
        else:
            self.full_report = report_data.get('full_report', {})
    
    def _get_field(self, field_name: str, default: Any = "") -> Any:
        """Get field from data or full_report, with mapping for new structure."""
        # Mapping from old field names to new field names
        field_mapping = {
            'summary': 'executive_summary',
            'key_highlights': 'key_strategies',
            'strategy': 'positioning',
            'marketing_plan': 'key_strategies',
            'revenue_model': 'budget_focus',
            'risk_analysis': 'growth_plan',
            'competitor_insights': 'positioning',
            'next_30_days': 'next_steps',
            'swot': 'swot_analysis'
        }
        
        # 1. Check direct field in data
        if field_name in self.data and self.data[field_name]:
            return self.data[field_name]
        
        # 2. Check new field name in full_report
        new_field = field_mapping.get(field_name)
        if new_field and new_field in self.full_report and self.full_report[new_field]:
            val = self.full_report[new_field]
            return val
            
        # 3. Check old field name in full_report
        if field_name in self.full_report and self.full_report[field_name]:
            return self.full_report[field_name]
            
        return default
    
    def _stringify_field(self, val: Any) -> str:
        """Convert list or dict to string for display."""
        if not val:
            return ""
        if isinstance(val, list):
            return "\n".join([f"- {s}" for s in val])
        if isinstance(val, dict):
            return "\n".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in val.items()])
        return str(val)

    def export_docx(self) -> bytes:
        """
        Generate DOCX file from report data.
        
        Returns:
            Bytes of the DOCX file
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Business Strategy Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle info
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.company_name:
            subtitle.add_run(f"Company: {self.company_name}\n").bold = True
        subtitle.add_run(f"Business Type: {self.business_type}\n").bold = True
        if self.location:
            subtitle.add_run(f"Location: {self.location}\n").bold = True
        subtitle.add_run(f"Target Audience: {self.target_audience}\n").bold = True
        subtitle.add_run(f"Budget: {self.budget}\n").bold = True
        subtitle.add_run(f"Generated: {self._format_date(self.created_at)}").bold = True
        
        doc.add_paragraph()
        
        # Executive Summary
        summary = self._stringify_field(self._get_field('summary'))
        if summary:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(summary)
            doc.add_paragraph()
        
        # USP
        if self.usp:
            doc.add_heading('Unique Selling Proposition', level=1)
            doc.add_paragraph(self.usp)
            doc.add_paragraph()
        
        # SWOT Analysis
        swot = self._get_field('swot')
        if swot and isinstance(swot, dict):
            doc.add_heading('SWOT Analysis', level=1)
            for key in ['strengths', 'weaknesses', 'opportunities', 'threats']:
                if key in swot:
                    doc.add_heading(key.title(), level=2)
                    for item in swot[key]:
                        doc.add_paragraph(item, style='List Bullet')
            doc.add_paragraph()
        
        # Strategy
        strategy = self._stringify_field(self._get_field('strategy'))
        if strategy:
            doc.add_heading('Business Strategy & Positioning', level=1)
            self._add_formatted_content(doc, strategy)
            doc.add_paragraph()
        
        # Marketing Plan
        marketing = self._stringify_field(self._get_field('marketing_plan'))
        if marketing:
            doc.add_heading('Strategic Actions', level=1)
            self._add_formatted_content(doc, marketing)
            doc.add_paragraph()
        
        # Roadmap (Next 30 Days / Next Steps)
        roadmap = self._stringify_field(self._get_field('next_30_days'))
        if roadmap:
            doc.add_heading('Action Roadmap', level=1)
            self._add_formatted_content(doc, roadmap)
            doc.add_paragraph()

        # Revenue Model
        revenue = self._stringify_field(self._get_field('revenue_model'))
        if revenue:
            doc.add_heading('Budget & Revenue Model', level=1)
            self._add_formatted_content(doc, revenue)
            doc.add_paragraph()
        
        # Save to bytes
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)
        return file_buffer.getvalue()
    
    def _add_formatted_content(self, doc: Document, content: str):
        """Add content with basic formatting (headers, bold, lists)."""
        if not content:
            return
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for markdown headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # Check for list items
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Check for numbered list
            elif line[0].isdigit() and '. ' in line[:3]:
                doc.add_paragraph(line, style='List Number')
            # Regular paragraph with bold detection
            else:
                # Process bold markers
                paragraph = doc.add_paragraph()
                parts = self._split_bold(line)
                for part in parts:
                    if part.get('bold'):
                        run = paragraph.add_run(part['text'])
                        run.bold = True
                    else:
                        paragraph.add_run(part['text'])
    
    def _split_bold(self, text: str) -> list:
        """Split text by ** markers for bold formatting."""
        import re
        parts = []
        last_end = 0
        
        for match in re.finditer(r'\*\*(.+?)\*\*', text):
            # Add text before match
            if match.start() > last_end:
                parts.append({'text': text[last_end:match.start()], 'bold': False})
            # Add bold text
            parts.append({'text': match.group(1), 'bold': True})
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            parts.append({'text': text[last_end:], 'bold': False})
        
        if not parts:
            parts.append({'text': text, 'bold': False})
        
        return parts
    
    def export_pdf(self) -> bytes:
        """
        Generate a professional PDF file from report data.
        """
        file_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            file_buffer,
            pagesize=LETTER,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Define high-end professional colors
        primary_color = colors.HexColor('#0F172A')  # Dark Slate
        secondary_color = colors.HexColor('#3B82F6')  # Modern Blue
        accent_color = colors.HexColor('#64748B')  # Slate Gray
        bg_light = colors.HexColor('#F8FAFC')
        
        # Custom styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=primary_color,
            spaceAfter=8,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=secondary_color,
            spaceAfter=24,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold',
            letterSpacing=1
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=primary_color,
            spaceBefore=20,
            spaceAfter=12,
            borderPadding=(5, 0, 5, 0),
            borderWidth=0,
            borderColor=primary_color,
            fontName='Helvetica-Bold'
        )
        
        # Draw a line after section headers
        def draw_section_header(title):
            story.append(Paragraph(title.upper(), section_header_style))
            story.append(Spacer(1, 2))
            # Table-based line for better control
            line = Table([['']], colWidths=[510], rowHeights=[2])
            line.setStyle(TableStyle([
                ('LINEBELOW', (0,0), (-1,-1), 2, secondary_color),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(line)
            story.append(Spacer(1, 10))

        normal_style = ParagraphStyle(
            'ModernNormal',
            parent=styles['Normal'],
            fontSize=10.5,
            textColor=colors.HexColor('#334155'),
            leading=15,
            spaceAfter=10,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        highlight_box_style = ParagraphStyle(
            'HighlightBox',
            parent=normal_style,
            fontSize=10,
            textColor=primary_color,
            backColor=bg_light,
            borderPadding=10,
            borderRadius=5,
            leftIndent=10,
            rightIndent=10,
            fontName='Helvetica-Oblique'
        )

        # 1. HEADER SECTION
        story.append(Paragraph("STRATEGY REPORT", subtitle_style))
        title_text = self.company_name if self.company_name else self.business_type
        story.append(Paragraph(f"{title_text}", title_style))
        
        # Info Table
        info_data = [
            [Paragraph(f"<b>TYPE:</b> {self.business_type}", normal_style), 
             Paragraph(f"<b>BUDGET:</b> {self.budget}", normal_style)],
            [Paragraph(f"<b>TARGET:</b> {self.target_audience}", normal_style),
             Paragraph(f"<b>LOCATION:</b> {self.location if self.location else 'Not specified'}", normal_style)],
            [Paragraph(f"<b>DATE:</b> {self._format_date(self.created_at)}", normal_style),
             Paragraph(f"<b>STRATEGY ID:</b> #{id(self.data) % 10000:04d}", normal_style)]
        ]
        info_table = Table(info_data, colWidths=[255, 255])
        info_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.lightgrey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 30))

        # 2. EXECUTIVE SUMMARY
        summary = self._stringify_field(self._get_field('summary'))
        if summary:
            draw_section_header('Executive Summary')
            story.append(Paragraph(summary, normal_style))
            story.append(Spacer(1, 15))

        # USP
        if self.usp:
            draw_section_header('Unique Selling Proposition')
            story.append(Paragraph(self.usp, highlight_box_style))
            story.append(Spacer(1, 15))

        # SWOT Analysis
        swot = self._get_field('swot')
        if swot and isinstance(swot, dict):
            draw_section_header('SWOT Analysis')
            swot_data = []
            
            # Helper to format SWOT list
            def fmt_swot(items):
                if not items: return "N/A"
                return "<br/>".join([f"• {item}" for item in items])

            swot_data = [
                [Paragraph("<b>STRENGTHS</b>", normal_style), Paragraph("<b>WEAKNESSES</b>", normal_style)],
                [Paragraph(fmt_swot(swot.get('strengths')), normal_style), Paragraph(fmt_swot(swot.get('weaknesses')), normal_style)],
                [Paragraph("<b>OPPORTUNITIES</b>", normal_style), Paragraph("<b>THREATS</b>", normal_style)],
                [Paragraph(fmt_swot(swot.get('opportunities')), normal_style), Paragraph(fmt_swot(swot.get('threats')), normal_style)]
            ]
            
            swot_table = Table(swot_data, colWidths=[255, 255])
            swot_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BACKGROUND', (0,0), (1,0), bg_light),
                ('BACKGROUND', (0,2), (1,2), bg_light),
                ('PADDING', (0,0), (-1,-1), 10),
            ]))
            story.append(swot_table)
            story.append(Spacer(1, 20))

        # 4. CORE SECTIONS
        sections = [
            ('Business Strategy & Positioning', 'strategy'),
            ('Strategic Actions', 'marketing_plan'),
            ('Budget & Revenue Model', 'revenue_model'),
            ('Action Roadmap', 'next_30_days')
        ]

        for title, field in sections:
            val = self._get_field(field)
            content = self._stringify_field(val)
            if content:
                draw_section_header(title)
                self._add_pdf_content(story, content, normal_style)
                story.append(Spacer(1, 15))

        # Build PDF
        doc.build(story)
        file_buffer.seek(0)
        return file_buffer.getvalue()

    def _add_pdf_content(self, story: list, content: str, base_style: ParagraphStyle):
        """Add formatted content to PDF story with bold support."""
        if not content:
            return
        
        # Define styles for markdown headers
        h2_style = ParagraphStyle('H2', parent=base_style, fontSize=12, fontName='Helvetica-Bold', textColor=colors.HexColor('#1E293B'), spaceBefore=10, spaceAfter=5)
        h3_style = ParagraphStyle('H3', parent=base_style, fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#475569'), spaceBefore=8, spaceAfter=4)
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Convert markdown bold **text** to <b>text</b>
            import re
            line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            
            # Check for headers
            if line.startswith('### '):
                story.append(Paragraph(line[4:], h3_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], h2_style))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], h2_style))
            # List items
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"&bull; {line[2:]}", base_style))
            elif line[0].isdigit() and '. ' in line[:4]:
                story.append(Paragraph(line, base_style))
            # Regular text
            else:
                story.append(Paragraph(line, base_style))
    
    def _format_date(self, date_string: str) -> str:
        """Format date string for display."""
        try:
            dt = datetime.fromisoformat(date_string.replace('Z', '+00:00'))
            return dt.strftime('%B %d, %Y at %H:%M')
        except (ValueError, TypeError):
            return str(date_string)


def export_report_as_docx(report_data: Dict[str, Any]) -> bytes:
    """
    Convenience function to export report as DOCX.
    
    Args:
        report_data: Dictionary containing report fields
        
    Returns:
        DOCX file as bytes
    """
    exporter = StrategyReportExporter(report_data)
    return exporter.export_docx()


def export_report_as_pdf(report_data: Dict[str, Any]) -> bytes:
    """
    Convenience function to export report as PDF.
    
    Args:
        report_data: Dictionary containing report fields
        
    Returns:
        PDF file as bytes
    """
    exporter = StrategyReportExporter(report_data)
    return exporter.export_pdf()

